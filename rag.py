import os
import json
import asyncio
import logging
import requests
import numpy as np
from datetime import datetime
from typing import List, Dict, Optional, Any, Tuple
from dataclasses import dataclass
from pathlib import Path
import tempfile
import time
import re
import hashlib
import email
import PyPDF2
import docx
from urllib.parse import urlparse, unquote

# FastAPI and Pydantic
from fastapi import FastAPI, HTTPException, Depends, Security, BackgroundTasks
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field

# Document Processing
from email.mime.text import MIMEText
from email.parser import Parser

# Vector Database and Embeddings
from sentence_transformers import SentenceTransformer
from pinecone import Pinecone, ServerlessSpec
from langchain_community.embeddings import SentenceTransformerEmbeddings

# LangChain Components
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_core.messages import BaseMessage, HumanMessage, AIMessage, SystemMessage
from langchain_groq import ChatGroq

# LangGraph for Multi-Agent System
from typing_extensions import TypedDict
from langgraph.graph import END, START, StateGraph
from langgraph.graph.message import add_messages
from typing import Annotated

# Web server
import uvicorn

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables from .env file
from pathlib import Path
env_file = Path(".env")
if env_file.exists():
    with open(env_file, 'r') as f:
        for line in f:
            line = line.strip()
            if '=' in line and not line.startswith('#'):
                key, value = line.split('=', 1)
                os.environ[key] = value

# Configuration
class Config:
    GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
    PINECONE_API_KEY = os.getenv("PINECONE_API_KEY", "")
    BEARER_TOKEN = "ae717b767b095a88fb3ec49728a42e27cb58162f1e6d67696bfe335a5c3c87c2"
    EMBEDDING_MODEL = "all-MiniLM-L6-v2"
    MAX_CHUNK_SIZE = 512
    CHUNK_OVERLAP = 100
    BASE_SIMILARITY_THRESHOLD = 0.6
    ADAPTIVE_THRESHOLD_RANGE = (0.5, 0.8)
    MIN_RELEVANT_CHUNKS = 3
    MAX_RELEVANT_CHUNKS = 8
    RERANK_TOP_K = 12
    PINECONE_INDEX_NAME = "insurance-documents"
    PINECONE_ENVIRONMENT = "us-west1-gcp"

config = Config()

# Initialize LLM
llm = ChatGroq(
    api_key=config.GROQ_API_KEY,
    model="llama3-8b-8192",
    temperature=0.1
)

# Initialize embedding model
embedding_model = SentenceTransformer(config.EMBEDDING_MODEL)
security = HTTPBearer()

# Pydantic Models
class QueryRequest(BaseModel):
    documents: str = Field(..., description="URL to the document blob (can be comma-separated for multiple documents)")
    questions: List[str] = Field(..., description="List of questions to answer")

class QueryResponse(BaseModel):
    answers: List[str] = Field(..., description="List of answers corresponding to questions")

# Enhanced Document Processing Classes
class AdvancedDocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.MAX_CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
            keep_separator=True
        )
        # Add document cache
        self.document_cache = {}

    def download_document(self, url: str) -> bytes:
        """Download document from URL with retry logic"""
        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = requests.get(url, timeout=30)
                response.raise_for_status()
                return response.content
            except requests.RequestException as e:
                if attempt == max_retries - 1:
                    raise HTTPException(status_code=400, detail=f"Failed to download document: {str(e)}")
                time.sleep(1)

    def extract_text_from_pdf(self, content: bytes) -> str:
        """Enhanced PDF text extraction with proper file handle management"""
        try:
            # Create temporary file but don't use context manager for the file creation
            tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            try:
                tmp_file.write(content)
                tmp_file.flush()
                tmp_file.close()  # Important: Close the file handle before reading
                
                # Now read the PDF
                text = ""
                with open(tmp_file.name, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            # Clean and format text
                            page_text = self._clean_text(page_text)
                            if page_text.strip():
                                text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                        except Exception as page_error:
                            logger.warning(f"Failed to extract page {page_num + 1}: {str(page_error)}")
                            continue
                
                return text
                
            finally:
                # Ensure file is deleted even if an exception occurs
                try:
                    if os.path.exists(tmp_file.name):
                        os.unlink(tmp_file.name)
                except Exception as cleanup_error:
                    logger.warning(f"Failed to cleanup temporary file: {str(cleanup_error)}")
                    
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to extract PDF text: {str(e)}")

    def extract_text_from_docx(self, content: bytes) -> str:
        """Enhanced DOCX text extraction"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(content)
                tmp_file.flush()

                doc = docx.Document(tmp_file.name)
                text = ""

                # Extract paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n"

                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            text += row_text + "\n"

                os.unlink(tmp_file.name)
                return self._clean_text(text)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to extract DOCX text: {str(e)}")
    
    def extract_text_from_email(self, content: bytes) -> str:
        """Enhanced email text extraction"""
        try:
            email_content = content.decode('utf-8')
            msg = email.message_from_string(email_content)

            text = f"Subject: {msg.get('Subject', 'No Subject')}\n"
            text += f"From: {msg.get('From', 'Unknown')}\n"
            text += f"Date: {msg.get('Date', 'Unknown')}\n\n"

            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        payload = part.get_payload(decode=True)
                        if payload:
                            text += payload.decode('utf-8', errors='ignore')
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    text += payload.decode('utf-8', errors='ignore')

            return self._clean_text(text)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to extract email text: {str(e)}")

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        # Remove special characters that might interfere
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        return text.strip()

    def process_document(self, url: str) -> List[Document]:
        """Enhanced document processing with caching"""
        # Check cache first
        if url in self.document_cache:
            logger.info(f"Using cached document for {url}")
            return self.document_cache[url]
        
        content = self.download_document(url)
        
        # Parse URL and extract the actual file path
        parsed_url = urlparse(url)
        file_path = unquote(parsed_url.path)
        file_path_lower = file_path.lower()
        
        # Determine file type from the actual file path
        if file_path_lower.endswith('.pdf') or content[:4] == b'%PDF':
            text = self.extract_text_from_pdf(content)
            doc_type = "PDF"
        elif file_path_lower.endswith(('.docx', '.doc')):
            text = self.extract_text_from_docx(content)
            doc_type = "DOCX"
        elif file_path_lower.endswith(('.eml', '.msg')):
            text = self.extract_text_from_email(content)
            doc_type = "EMAIL"
        else:
            try:
                text = content.decode('utf-8', errors='ignore')
                doc_type = "TEXT"
            except:
                raise HTTPException(status_code=400, detail=f"Unsupported document format. URL path: {file_path}")
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No text content found in document")
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        documents = []
        
        for i, chunk in enumerate(chunks):
            if chunk.strip():
                metadata = {
                    "source": url,
                    "chunk_id": i,
                    "doc_type": doc_type,
                    "chunk_length": len(chunk),
                    "word_count": len(chunk.split())
                }
                documents.append(Document(page_content=chunk, metadata=metadata))
        
        # Cache the processed documents
        self.document_cache[url] = documents
        logger.info(f"Cached {len(documents)} document chunks for {url}")
        
        return documents

    def debug_url_parsing(self, url: str):
        """Debug URL parsing"""
        parsed_url = urlparse(url)
        file_path = unquote(parsed_url.path)
        print(f"Original URL: {url}")
        print(f"Parsed path: {file_path}")
        print(f"File type detection: {file_path.lower().endswith('.pdf')}")

# Insurance-Specific Processing Classes
@dataclass
class InsuranceEntity:
    entity_type: str
    value: str
    context: str
    confidence: float

class InsuranceDocumentProcessor:
    def __init__(self):
        self.insurance_patterns = {
            'policy_numbers': r'(?i)policy\s+(?:no\.?|number)\s*:?\s*([A-Z0-9/-]+)',
            'claim_numbers': r'(?i)claim\s+(?:no\.?|number)\s*:?\s*([A-Z0-9/-]+)',
            'amounts': r'(?:₹|INR|Rs\.?)\s*([0-9,]+(?:\.[0-9]{2})?)|([0-9,]+(?:\.[0-9]{2})?)\s*(?:₹|INR|Rs\.?)',
            'percentages': r'(\d+(?:\.\d+)?)\s*%',
            'waiting_periods': r'(\d+)\s*(?:days?|months?|years?)\s*(?:waiting\s*period|wait)',
            'grace_periods': r'(\d+)\s*(?:days?|months?)\s*(?:grace\s*period)',
            'coverage_limits': r'(?i)(?:sum\s+insured|coverage|limit).*?(?:₹|INR|Rs\.?)\s*([0-9,]+(?:\.[0-9]{2})?)',
            'deductibles': r'(?i)deductible.*?(?:₹|INR|Rs\.?)\s*([0-9,]+(?:\.[0-9]{2})?)',
            'exclusions': r'(?i)(?:excluded?|not\s+covered?|exception)',
            'inclusions': r'(?i)(?:included?|covered?|benefit)',
            'medical_conditions': r'(?i)(diabetes|cancer|heart\s+disease|hypertension|kidney|liver|stroke|surgery|maternity|pregnancy)',
            'age_limits': r'(?i)(?:age|years?)\s*(?:limit|between|from|to)\s*(\d+)(?:\s*(?:to|-)?\s*(\d+))?',
            'renewal_terms': r'(?i)renewal.*?(\d+)\s*(?:days?|months?|years?)',
            'co_payment': r'(?i)co-?pay(?:ment)?.*?(\d+(?:\.\d+)?)\s*%'
        }
        
        self.section_headers = [
            r'(?i)^(?:section|clause|article|paragraph|part)\s+(\d+(?:\.\d+)*)',
            r'(?i)^([A-Z][A-Z\s&]+):',  # ALL CAPS headers
            r'(?i)^(\d+\.\s+[A-Z][A-Za-z\s]+)',  # Numbered sections
        ]
    
    def extract_structured_entities(self, text: str) -> List[InsuranceEntity]:
        """Extract insurance-specific structured information"""
        entities = []
        
        for entity_type, pattern in self.insurance_patterns.items():
            matches = re.finditer(pattern, text)
            for match in matches:
                # Get surrounding context (50 chars before and after)
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                context = text[start:end].strip()
                
                # Extract the actual value
                value = match.group(1) if match.groups() else match.group(0)
                
                # Calculate confidence based on context quality
                confidence = self._calculate_entity_confidence(entity_type, value, context)
                
                entities.append(InsuranceEntity(
                    entity_type=entity_type,
                    value=value.strip(),
                    context=context,
                    confidence=confidence
                ))
        
        return entities
    
    def _calculate_entity_confidence(self, entity_type: str, value: str, context: str) -> float:
        """Calculate confidence score for extracted entity"""
        confidence = 0.5  # Base confidence
        
        # Boost confidence for specific patterns
        confidence_boosters = {
            'amounts': ['sum insured', 'premium', 'claim amount', 'benefit'],
            'waiting_periods': ['waiting period', 'wait', 'after'],
            'grace_periods': ['grace period', 'due date', 'payment'],
            'medical_conditions': ['covered', 'excluded', 'treatment', 'diagnosis']
        }
        
        if entity_type in confidence_boosters:
            for booster in confidence_boosters[entity_type]:
                if booster.lower() in context.lower():
                    confidence += 0.15
        
        return min(confidence, 1.0)
    
    def create_enhanced_chunks(self, documents: List[Document]) -> List[Document]:
        """Create enhanced chunks with extracted entities"""
        enhanced_docs = []
        
        for doc in documents:
            entities = self.extract_structured_entities(doc.page_content)
            
            # Create entity summary for metadata
            entity_summary = {}
            for entity in entities:
                if entity.entity_type not in entity_summary:
                    entity_summary[entity.entity_type] = []
                entity_summary[entity.entity_type].append(entity.value)
            
            # Enhanced metadata
            enhanced_metadata = doc.metadata.copy()
            enhanced_metadata.update({
                'entities': entity_summary,
                'entity_count': len(entities),
                'has_amounts': bool(entity_summary.get('amounts')),
                'has_waiting_periods': bool(entity_summary.get('waiting_periods')),
                'has_medical_conditions': bool(entity_summary.get('medical_conditions')),
                'structure_richness': len(entity_summary) / len(self.insurance_patterns) if self.insurance_patterns else 0
            })
            
            enhanced_docs.append(Document(
                page_content=doc.page_content,
                metadata=enhanced_metadata
            ))
        
        return enhanced_docs

class InsuranceQueryEnhancer:
    def __init__(self, llm):
        self.llm = llm
        self.insurance_synonyms = {
            'premium': ['payment', 'installment', 'contribution'],
            'coverage': ['benefit', 'protection', 'indemnity'],
            'deductible': ['excess', 'co-payment', 'out-of-pocket'],
            'exclusion': ['exception', 'limitation', 'restriction'],
            'waiting period': ['qualification period', 'probation period'],
            'grace period': ['extra time', 'extension period'],
            'claim': ['settlement', 'compensation', 'reimbursement'],
            'policy': ['contract', 'agreement', 'plan'],
            'maternity': ['pregnancy', 'childbirth', 'delivery'],
            'pre-existing': ['prior condition', 'existing illness']
        }
    
    def expand_insurance_query(self, query: str) -> str:
        """Expand query with insurance domain synonyms"""
        expanded_terms = [query]
        
        query_lower = query.lower()
        for main_term, synonyms in self.insurance_synonyms.items():
            if main_term in query_lower:
                expanded_terms.extend([f"({main_term}|{syn})" for syn in synonyms[:2]])  # Limit to 2 synonyms
        
        return ' '.join(expanded_terms)

class InsuranceEnsemble:
    def __init__(self):
        self.models = [
            ("llama3-8b", ChatGroq(model="llama3-8b-8192", temperature=0.1)),
        ]
        
    def ensemble_answer_insurance(self, context: str, query: str, entities: Dict) -> Tuple[str, float]:
        """Generate ensemble answer with insurance domain focus"""
        answers = []
        confidences = []
        
        # Insurance-specific system prompt
        insurance_system_prompt = """You are an expert insurance policy analyst with deep knowledge of Indian insurance regulations and terminology.

Your expertise includes:
- Policy terms, conditions, and exclusions
- Premium calculations and payment terms
- Claim procedures and settlement processes
- Waiting periods and grace periods  
- Coverage limits and deductibles
- Pre-existing disease clauses
- Maternity and specific treatment coverage

Instructions:
1. Provide precise, factual answers based only on the document context
2. Include specific amounts, percentages, and time periods when mentioned
3. Clearly state if information is subject to conditions or exclusions
4. Use proper insurance terminology
5. If information is incomplete, specify what details are missing"""
        
        for model_name, model in self.models:
            try:
                # Create model-specific prompt with entity context
                entity_info = self._format_entities_for_prompt(entities)
                
                user_prompt = f"""
Context Document: {context}

Key Entities Identified: {entity_info}

Insurance Question: {query}

Please provide a comprehensive answer focusing on the specific insurance terms and conditions mentioned in the document.
"""
                
                messages = [
                    SystemMessage(content=insurance_system_prompt),
                    HumanMessage(content=user_prompt)
                ]
                
                response = model.invoke(messages)
                answers.append((model_name, response.content))
                
                # Calculate confidence based on answer specificity
                confidence = self._calculate_insurance_answer_confidence(response.content, entities)
                confidences.append(confidence)
                
            except Exception as e:
                logger.warning(f"Model {model_name} failed: {e}")
                continue
        
        if not answers:
            return "Unable to generate answer from ensemble models", 0.0
        
        # Select best answer
        best_idx = np.argmax(confidences) if confidences else 0
        best_confidence = confidences[best_idx] if confidences else 0.0
        
        return answers[best_idx][1], best_confidence
    
    def _format_entities_for_prompt(self, entities: Dict) -> str:
        """Format entities for inclusion in prompt"""
        if not entities:
            return "No specific entities identified"
        
        formatted = []
        for entity_type, values in entities.items():
            if values:
                formatted.append(f"{entity_type.replace('_', ' ').title()}: {', '.join(values[:3])}")
        
        return "; ".join(formatted)
    
    def _calculate_insurance_answer_confidence(self, answer: str, entities: Dict) -> float:
        """Calculate confidence score for insurance-specific answers"""
        confidence = 0.5
        
        # Boost for specific insurance terms
        insurance_indicators = [
            'waiting period', 'grace period', 'premium', 'coverage', 'exclusion',
            'deductible', 'co-payment', 'sum insured', 'claim', 'benefit'
        ]
        
        for indicator in insurance_indicators:
            if indicator.lower() in answer.lower():
                confidence += 0.08
        
        # Boost for numerical specificity
        if re.search(r'\d+\s*(?:days?|months?|years?|%|₹|INR)', answer):
            confidence += 0.15
        
        # Boost for mentioning conditions/limitations
        condition_words = ['subject to', 'provided that', 'except', 'unless', 'condition']
        for condition in condition_words:
            if condition.lower() in answer.lower():
                confidence += 0.05
        
        return min(confidence, 1.0)

# Optimized Pinecone Vector Store for Speed
class PineconeInsuranceVectorStore:
    def __init__(self, api_key: str, index_name: str = "insurance-documents"):
        # Initialize Pinecone with new API
        self.pc = Pinecone(api_key=api_key)
        self.index_name = index_name
        self.dimension = 384  # SentenceTransformer dimension
        
        # Create index if it doesn't exist
        existing_indexes = [index.name for index in self.pc.list_indexes()]
        if index_name not in existing_indexes:
            self.pc.create_index(
                name=index_name,
                dimension=self.dimension,
                metric="cosine",
                spec=ServerlessSpec(
                    cloud="aws",
                    region="us-east-1"
                )
            )
        
        self.index = self.pc.Index(index_name)
        self.embeddings = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
        
        # Cache to track processed documents and their vector IDs
        self.processed_docs = set()
        self.document_vectors = {}  # doc_hash -> list of vector_ids
    
    def document_exists(self, doc_url: str) -> bool:
        """Check if document is already indexed"""
        doc_hash = hashlib.md5(doc_url.encode()).hexdigest()
        return doc_hash in self.processed_docs
    
    def get_document_hash(self, doc_url: str) -> str:
        """Get document hash for tracking"""
        return hashlib.md5(doc_url.encode()).hexdigest()
    
    async def delete_document_embeddings(self, doc_url: str):
        """Delete all embeddings for a specific document"""
        doc_hash = self.get_document_hash(doc_url)
        
        try:
            # Get all vector IDs for this document
            # Since Pinecone doesn't support listing by prefix easily, we'll use a query approach
            # We'll track vector IDs during insertion for easier deletion
            if hasattr(self, 'document_vectors'):
                vector_ids = self.document_vectors.get(doc_hash, [])
                if vector_ids:
                    # Delete in batches
                    batch_size = 100
                    for i in range(0, len(vector_ids), batch_size):
                        batch_ids = vector_ids[i:i + batch_size]
                        self.index.delete(ids=batch_ids)
                    
                    # Remove from tracking
                    del self.document_vectors[doc_hash]
                    logger.info(f"Deleted {len(vector_ids)} embeddings for document {doc_url}")
            
            # Remove from processed docs
            if doc_hash in self.processed_docs:
                self.processed_docs.remove(doc_hash)
                
        except Exception as e:
            logger.error(f"Error deleting embeddings for {doc_url}: {str(e)}")
    
    async def delete_multiple_documents(self, doc_urls: List[str]):
        """Delete embeddings for multiple documents"""
        for doc_url in doc_urls:
            await self.delete_document_embeddings(doc_url)
    
    async def add_documents_async(self, documents: List[Document], doc_url: str):
        """Add documents to Pinecone with async batching"""
        doc_hash = hashlib.md5(doc_url.encode()).hexdigest()
        
        if self.document_exists(doc_url):
            logger.info(f"Document {doc_url} already indexed, skipping")
            return
        
        # Process documents with insurance enhancements
        insurance_processor = InsuranceDocumentProcessor()
        enhanced_docs = insurance_processor.create_enhanced_chunks(documents)
        
        # Prepare vectors for batch upload
        vectors_to_upsert = []
        vector_ids = []  # Track vector IDs for this document
        batch_size = 100
        
        for i, doc in enumerate(enhanced_docs):
            # Generate embedding
            embedding = self.embeddings.embed_query(doc.page_content)
            
            # Create unique ID
            vector_id = f"{doc_hash}_{i}"
            vector_ids.append(vector_id)
            
            # Prepare metadata
            metadata = {
                "text": doc.page_content[:1000],  # Pinecone metadata limit
                "source": doc_url,
                "chunk_id": i,
                **{k: str(v)[:100] for k, v in doc.metadata.items() if k != 'entities'}
            }
            
            vectors_to_upsert.append({
                "id": vector_id,
                "values": embedding,
                "metadata": metadata
            })
            
            # Batch upload every 100 vectors
            if len(vectors_to_upsert) >= batch_size:
                self.index.upsert(vectors=vectors_to_upsert)
                vectors_to_upsert = []
                await asyncio.sleep(0.1)  # Small delay to avoid rate limits
        
        # Upload remaining vectors
        if vectors_to_upsert:
            self.index.upsert(vectors=vectors_to_upsert)
        
        # Track vector IDs and mark document as processed
        self.document_vectors[doc_hash] = vector_ids
        self.processed_docs.add(doc_hash)
        logger.info(f"Successfully indexed {len(enhanced_docs)} chunks for {doc_url}")
        logger.info(f"Tracked {len(vector_ids)} vector IDs for future deletion")
    
    async def similarity_search_async(self, query: str, top_k: int = 8) -> List[Tuple[Document, float]]:
        """Fast similarity search using Pinecone"""
        # Generate query embedding
        query_embedding = self.embeddings.embed_query(query)
        
        # Search Pinecone
        results = self.index.query(
            vector=query_embedding,
            top_k=top_k,
            include_metadata=True
        )
        
        # Convert results to Document format
        retrieved_docs = []
        for match in results.matches:
            doc = Document(
                page_content=match.metadata.get("text", ""),
                metadata=match.metadata
            )
            retrieved_docs.append((doc, float(match.score)))
        
        return retrieved_docs

# Query Classification for Adaptive Processing
class QueryClassifier:
    @staticmethod
    def classify_question_type(question: str) -> str:
        """Classify question type for adaptive processing"""
        question_lower = question.lower()
        
        # Specific question patterns
        specific_patterns = [
            r'\b(what is|define|definition of|meaning of)\b',
            r'\b(how much|how many|what amount|specific)\b',
            r'\b(when|what date|what time|deadline)\b',
            r'\b(where|location|address|place)\b',
            r'\b(who|which person|contact)\b',
            r'\b(clause|section|article|paragraph)\b',
            r'\b(percentage|rate|fee|cost|price)\b'
        ]
        
        # General question patterns
        general_patterns = [
            r'\b(explain|describe|tell me about|overview)\b',
            r'\b(how does|how can|what are the ways)\b',
            r'\b(benefits|advantages|features)\b',
            r'\b(process|procedure|steps)\b'
        ]
        
        for pattern in specific_patterns:
            if re.search(pattern, question_lower):
                return "specific"
        
        for pattern in general_patterns:
            if re.search(pattern, question_lower):
                return "general"
        
        return "general"  # Default to general
    
    @staticmethod
    def extract_key_entities(question: str) -> List[str]:
        """Extract key entities from question for better retrieval"""
        # Simple entity extraction (can be enhanced with NER)
        entities = []
        
        # Extract potential entities (capitalized words, numbers, etc.)
        words = question.split()
        for word in words:
            cleaned_word = re.sub(r'[^\w]', '', word)
            if (cleaned_word.isupper() or 
                cleaned_word.istitle() or 
                cleaned_word.isdigit() or 
                len(cleaned_word) > 8):  # Longer words might be important terms
                entities.append(cleaned_word.lower())
        
        return entities

# Enhanced Multi-Agent System State
class AgentState(TypedDict):
    messages: Annotated[List[BaseMessage], add_messages]
    query: str
    question_type: str
    key_entities: List[str]
    documents: List[Document]
    vector_store: PineconeInsuranceVectorStore
    retrieved_docs: List[Tuple[Document, float]]
    entities: Dict  # Added for insurance entities
    answer: str
    confidence: float
    method_used: str
    num_chunks_used: int

# Enhanced Agent Functions
def optimized_document_processor_agent(state: AgentState) -> AgentState:
    """Optimized document processing with Pinecone caching"""
    global shared_document_processor, pinecone_vector_store
    
    try:
        # Extract document URL
        documents_url = None
        for msg in state["messages"]:
            if hasattr(msg, 'content') and 'documents' in str(msg.content):
                json_match = re.search(r'\{.*\}', str(msg.content), re.DOTALL)
                if json_match:
                    data = json.loads(json_match.group())
                    documents_url = data.get('documents')
                    break
        
        if not documents_url:
            raise HTTPException(status_code=400, detail="No document URL provided")
        
        # Check if already processed
        if pinecone_vector_store.document_exists(documents_url):
            logger.info(f"Document already indexed: {documents_url}")
            state["vector_store"] = pinecone_vector_store
            state["documents"] = []  # No need to store in memory
        else:
            # Process and index new document
            documents = shared_document_processor.process_document(documents_url)
            
            # Upload to Pinecone asynchronously
            asyncio.create_task(
                pinecone_vector_store.add_documents_async(documents, documents_url)
            )
            
            state["vector_store"] = pinecone_vector_store
            state["documents"] = documents
            
        return state
        
    except Exception as e:
        logger.error(f"Document processing error: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Document processing failed: {str(e)}")

def document_processor_agent(state: AgentState) -> AgentState:
    """Enhanced document processing and indexing with shared processor"""
    global shared_document_processor, pinecone_vector_store
    
    try:
        # Extract documents URL from messages
        documents_url = None
        for msg in state["messages"]:
            if hasattr(msg, 'content') and 'documents' in str(msg.content):
                # Parse JSON from message content
                json_match = re.search(r'\{.*\}', str(msg.content), re.DOTALL)
                if json_match:
                    data = json.loads(json_match.group())
                    documents_url = data.get('documents')
                    break
        
        if not documents_url:
            documents_url = state.get("documents_url", "")
        
        if documents_url:
            # Check if already processed in Pinecone
            if pinecone_vector_store.document_exists(documents_url):
                logger.info(f"Document already indexed: {documents_url}")
                state["vector_store"] = pinecone_vector_store
                state["documents"] = []
            else:
                # Process and upload to Pinecone
                documents = shared_document_processor.process_document(documents_url)
                asyncio.create_task(
                    pinecone_vector_store.add_documents_async(documents, documents_url)
                )
                state["documents"] = documents
                state["vector_store"] = pinecone_vector_store
                logger.info(f"Processed {len(documents)} document chunks for Pinecone upload")
        else:
            state["documents"] = []
            state["vector_store"] = pinecone_vector_store
            
        return state
    except HTTPException:
        # Re-raise HTTP exceptions to be handled by FastAPI
        raise
    except Exception as e:
        logger.error(f"Document processing error: {str(e)}")
        # For non-HTTP exceptions, convert to HTTP exception
        raise HTTPException(status_code=400, detail=f"Document processing failed: {str(e)}")

def query_analysis_agent(state: AgentState) -> AgentState:
    """Analyze query for adaptive processing"""
    try:
        query = state["query"]
        
        # Classify question type
        question_type = QueryClassifier.classify_question_type(query)
        
        # Extract key entities
        key_entities = QueryClassifier.extract_key_entities(query)
        
        state["question_type"] = question_type
        state["key_entities"] = key_entities
        
        logger.info(f"Query analysis - Type: {question_type}, Entities: {key_entities}")
        return state
    except Exception as e:
        logger.error(f"Query analysis error: {str(e)}")
        state["question_type"] = "general"
        state["key_entities"] = []
        return state

async def adaptive_retrieval_agent(state: AgentState) -> AgentState:
    """Enhanced retrieval with Pinecone fast search"""
    query = state["query"]
    question_type = state["question_type"]
    vector_store = state["vector_store"]
    
    try:
        # Use Pinecone fast similarity search
        retrieved_docs = await vector_store.similarity_search_async(query, top_k=config.MAX_RELEVANT_CHUNKS)
        
        state["retrieved_docs"] = retrieved_docs
        state["entities"] = {}  # Entities will be extracted from metadata if available
        state["num_chunks_used"] = len(retrieved_docs)
        
        # Calculate confidence based on retrieval quality
        if retrieved_docs:
            scores = [score for _, score in retrieved_docs]
            avg_score = np.mean(scores)
            score_variance = np.var(scores)
            
            # Higher confidence for consistent high scores
            confidence = avg_score * (1 - score_variance * 0.5)
            state["confidence"] = max(0.0, min(1.0, confidence))
        else:
            state["confidence"] = 0.0
            
        logger.info(f"Retrieved {len(retrieved_docs)} chunks with confidence {state['confidence']:.3f}")
        return state
    except Exception as e:
        logger.error(f"Retrieval error: {str(e)}")
        state["retrieved_docs"] = []
        state["confidence"] = 0.0
        state["num_chunks_used"] = 0
        state["entities"] = {}
        return state

def adaptive_retrieval_agent_sync(state: AgentState) -> AgentState:
    """Enhanced retrieval with adaptive techniques and insurance-specific search"""
    query = state["query"]
    question_type = state["question_type"]
    vector_store = state["vector_store"]
    
    try:
        # Use enhanced insurance search if available
        if hasattr(vector_store, 'enhanced_insurance_search'):
            retrieved_docs, entities = vector_store.enhanced_insurance_search(query, question_type)
            state["entities"] = entities  # Store entities for later use
        else:
            # Fallback to adaptive similarity search
            retrieved_docs = vector_store.adaptive_similarity_search(query, question_type)
            state["entities"] = {}
        
        state["retrieved_docs"] = retrieved_docs
        state["num_chunks_used"] = len(retrieved_docs)
        
        # Calculate confidence based on retrieval quality
        if retrieved_docs:
            scores = [score for _, score in retrieved_docs]
            avg_score = np.mean(scores)
            score_variance = np.var(scores)
            
            # Higher confidence for consistent high scores
            confidence = avg_score * (1 - score_variance * 0.5)
            state["confidence"] = max(0.0, min(1.0, confidence))
        else:
            state["confidence"] = 0.0
            
        logger.info(f"Retrieved {len(retrieved_docs)} chunks with confidence {state['confidence']:.3f}")
        return state
    except Exception as e:
        logger.error(f"Retrieval error: {str(e)}")
        state["retrieved_docs"] = []
        state["confidence"] = 0.0
        state["num_chunks_used"] = 0
        state["entities"] = {}
        return state

def enhanced_insurance_llm_agent(state: AgentState) -> AgentState:
    """Enhanced LLM agent specifically for insurance documents"""
    query = state["query"]
    retrieved_docs = state["retrieved_docs"]
    question_type = state["question_type"]
    key_entities = state["key_entities"]
    entities = state.get("entities", {})
    vector_store = state["vector_store"]
    
    try:
        # Lower the confidence threshold - the current threshold is too high
        min_confidence_threshold = 0.3  # Lowered from 0.5
        
        if retrieved_docs and (state["confidence"] > min_confidence_threshold or len(retrieved_docs) >= config.MIN_RELEVANT_CHUNKS):
            # Prepare context with chunk ranking
            context_parts = []
            for i, (doc, score) in enumerate(retrieved_docs, 1):
                context_parts.append(f"[Chunk {i} - Relevance: {score:.3f}]\n{doc.page_content}")
            
            context = "\n\n".join(context_parts)
            
            # Use ensemble approach for better accuracy if available
            if hasattr(vector_store, 'ensemble') and entities:
                answer, confidence = vector_store.ensemble.ensemble_answer_insurance(
                    context, query, entities
                )
                state["answer"] = answer
                state["confidence"] = confidence
                state["method_used"] = "ENHANCED_INSURANCE_ENSEMBLE"
            else:
                # Fallback to enhanced insurance-specific prompting
                # Insurance-specific system prompt
                insurance_system_prompt = """You are an expert insurance policy analyst with deep knowledge of Indian insurance regulations and terminology.

Your expertise includes:
- Policy terms, conditions, and exclusions
- Premium calculations and payment terms
- Claim procedures and settlement processes
- Waiting periods and grace periods  
- Coverage limits and deductibles
- Pre-existing disease clauses
- Maternity and specific treatment coverage

Instructions:
1. Provide precise, factual answers based only on the document context
2. Include specific amounts, percentages, and time periods when mentioned
3. Clearly state if information is subject to conditions or exclusions
4. Use proper insurance terminology
5. If information is incomplete, specify what details are missing"""

                # Enhanced user prompt with context awareness
                entity_info = _format_entities_for_prompt(entities) if entities else "No specific entities identified"
                
                user_prompt = f"""Document Context (ranked by relevance):
{context}

Question: {query}
Question Type: {question_type}
Key Terms: {', '.join(key_entities) if key_entities else 'None identified'}
Insurance Entities Found: {entity_info}

Instructions:
1. Answer the question using ONLY the information provided in the document context
2. Focus on insurance-specific terms, amounts, percentages, and time periods
3. If the context contains multiple relevant pieces of information, synthesize them coherently
4. If the context doesn't fully address the question, clearly state what information is missing
5. Include specific references to relevant chunks when citing information
6. Maintain accuracy and avoid making assumptions beyond the provided context

Please provide your response:"""

                messages = [
                    SystemMessage(content=insurance_system_prompt),
                    HumanMessage(content=user_prompt)
                ]
                
                response = llm.invoke(messages)
                state["answer"] = response.content
                state["method_used"] = "ENHANCED_INSURANCE_RAG"
            
        else:
            # Insufficient context - provide informative response
            state["answer"] = ("I apologize, but I couldn't find sufficient relevant information in the provided document to answer your question accurately. "
                             "The document may not contain the specific information you're looking for, or the question might require additional context not present in the current document.")
            state["method_used"] = "INSUFFICIENT_CONTEXT"
            state["confidence"] = 0.0
            
        return state
    except Exception as e:
        logger.error(f"Enhanced insurance LLM agent error: {str(e)}")
        state["answer"] = f"I encountered an error while processing your insurance question: {str(e)}"
        state["method_used"] = "ERROR"
        return state

def _format_entities_for_prompt(entities: Dict) -> str:
    """Format entities for inclusion in prompt"""
    if not entities:
        return "No specific entities identified"
    
    formatted = []
    for entity_type, values in entities.items():
        if values:
            formatted.append(f"{entity_type.replace('_', ' ').title()}: {', '.join(values[:3])}")
    
    return "; ".join(formatted)
    """Enhanced LLM agent with adaptive prompting"""
    query = state["query"]
    retrieved_docs = state["retrieved_docs"]
    question_type = state["question_type"]
    key_entities = state["key_entities"]
    
    try:
        # Lower the confidence threshold - the current threshold is too high
        min_confidence_threshold = 0.3  # Lowered from 0.5
        
        if retrieved_docs and (state["confidence"] > min_confidence_threshold or len(retrieved_docs) >= config.MIN_RELEVANT_CHUNKS):
            # Prepare context with chunk ranking
            context_parts = []
            for i, (doc, score) in enumerate(retrieved_docs, 1):
                context_parts.append(f"[Chunk {i} - Relevance: {score:.3f}]\n{doc.page_content}")
            
            context = "\n\n".join(context_parts)
            
            # Adaptive system prompt based on question type
            if question_type == "specific":
                system_prompt = """You are a precision document analyst specializing in extracting specific, factual information from insurance, legal, HR, and compliance documents.

Your mission is to provide EXACT, SPECIFIC answers with precise details:
- Extract specific numbers, dates, percentages, amounts, and terms
- Quote exact clauses, sections, or conditions when relevant
- If information is not explicitly stated, clearly indicate this
- Prioritize accuracy over completeness
- Always cite the specific source or section when possible

Focus on delivering concise, factual responses with specific details."""
            else:  # general questions
                system_prompt = """You are a comprehensive document analyst specializing in insurance, legal, HR, and compliance domains.

Your mission is to provide thorough, well-structured explanations:
- Provide comprehensive overviews and explanations
- Structure information logically with clear sections
- Include relevant context and background information
- Explain processes, procedures, and concepts clearly
- Connect related information across document sections
- Provide actionable insights when appropriate

Focus on delivering complete, well-organized responses that fully address the question."""

            # Enhanced user prompt with context awareness
            user_prompt = f"""Document Context (ranked by relevance):
{context}

Question: {query}
Question Type: {question_type}
Key Terms: {', '.join(key_entities) if key_entities else 'None identified'}

Instructions:
1. Answer the question using ONLY the information provided in the document context
2. If the context contains multiple relevant pieces of information, synthesize them coherently
3. If the context doesn't fully address the question, clearly state what information is missing
4. Include specific references to relevant chunks when citing information
5. Maintain accuracy and avoid making assumptions beyond the provided context

Please provide your response:"""

            messages = [
                SystemMessage(content=system_prompt),
                HumanMessage(content=user_prompt)
            ]
            
            response = llm.invoke(messages)
            state["answer"] = response.content
            state["method_used"] = "ADAPTIVE_RAG"
            
        else:
            # Insufficient context - provide informative response
            state["answer"] = ("I apologize, but I couldn't find sufficient relevant information in the provided document to answer your question accurately. "
                             "The document may not contain the specific information you're looking for, or the question might require additional context not present in the current document.")
            state["method_used"] = "INSUFFICIENT_CONTEXT"
            state["confidence"] = 0.0
            
        return state
    except Exception as e:
        logger.error(f"LLM agent error: {str(e)}")
        state["answer"] = f"I encountered an error while processing your question: {str(e)}"
        state["method_used"] = "ERROR"
        return state

# Build Enhanced Multi-Agent Graph
def build_optimized_agent_graph():
    """Build the optimized multi-agent workflow graph with Pinecone"""
    builder = StateGraph(AgentState)
    
    # Add nodes
    builder.add_node("optimized_document_processor", optimized_document_processor_agent)
    builder.add_node("query_analysis", query_analysis_agent)
    builder.add_node("adaptive_retrieval", adaptive_retrieval_agent_sync)
    builder.add_node("enhanced_insurance_llm", enhanced_insurance_llm_agent)
    
    # Add edges
    builder.add_edge(START, "optimized_document_processor")
    builder.add_edge("optimized_document_processor", "query_analysis")
    builder.add_edge("query_analysis", "adaptive_retrieval")
    builder.add_edge("adaptive_retrieval", "enhanced_insurance_llm")
    builder.add_edge("enhanced_insurance_llm", END)
    
    return builder.compile()

def build_enhanced_agent_graph():
    """Build the enhanced multi-agent workflow graph"""
    builder = StateGraph(AgentState)
    
    # Add nodes
    builder.add_node("document_processor", document_processor_agent)
    builder.add_node("query_analysis", query_analysis_agent)
    builder.add_node("adaptive_retrieval", adaptive_retrieval_agent_sync)
    builder.add_node("enhanced_insurance_llm", enhanced_insurance_llm_agent)
    
    # Add edges
    builder.add_edge(START, "document_processor")
    builder.add_edge("document_processor", "query_analysis")
    builder.add_edge("query_analysis", "adaptive_retrieval")
    builder.add_edge("adaptive_retrieval", "enhanced_insurance_llm")
    builder.add_edge("enhanced_insurance_llm", END)
    
    return builder.compile()

# Initialize both agent graphs
agent_graph = build_enhanced_agent_graph()
optimized_agent_graph = build_optimized_agent_graph()

# At the module level, create a shared document processor
shared_document_processor = AdvancedDocumentProcessor()

# Initialize Pinecone Vector Store
pinecone_vector_store = PineconeInsuranceVectorStore(config.PINECONE_API_KEY, config.PINECONE_INDEX_NAME)

# FastAPI App
app = FastAPI(
    title="Enhanced Adaptive RAG System",
    description="Advanced RAG system with adaptive retrieval and multi-agent architecture",
    version="2.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Authentication
async def verify_token(credentials: HTTPAuthorizationCredentials = Security(security)):
    if credentials.credentials != config.BEARER_TOKEN:
        raise HTTPException(status_code=401, detail="Invalid authentication token")
    return credentials.credentials

# API Endpoints
@app.post("/hackrx/run", response_model=QueryResponse)
async def optimized_query_retrieval(
    request: QueryRequest,
    background_tasks: BackgroundTasks,
    token: str = Depends(verify_token)
):
    """Optimized endpoint with Pinecone and parallel processing"""
    start_time = time.time()
    
    try:
        # Parse documents URLs (support multiple documents)
        documents_urls = [url.strip() for url in request.documents.split(',') if url.strip()]
        if not documents_urls:
            raise HTTPException(status_code=400, detail="No document URLs provided")
        
        logger.info(f"Processing {len(documents_urls)} documents and {len(request.questions)} questions")
        
        # Process all documents first
        for doc_url in documents_urls:
            if not pinecone_vector_store.document_exists(doc_url):
                logger.info(f"Processing new document: {doc_url}")
                documents = shared_document_processor.process_document(doc_url)
                await pinecone_vector_store.add_documents_async(documents, doc_url)
            else:
                logger.info(f"Document already indexed: {doc_url}")
        
        # Process all questions in parallel
        async def process_single_question(question: str) -> str:
            try:
                # Fast similarity search across all indexed documents
                retrieved_docs = await pinecone_vector_store.similarity_search_async(question)
                
                if not retrieved_docs:
                    return "No relevant information found in the documents."
                
                # Prepare context from top relevant chunks
                context = "\n\n".join([
                    f"[Chunk {i+1} - Relevance: {score:.3f}]\n{doc.page_content}"
                    for i, (doc, score) in enumerate(retrieved_docs[:6])
                ])
                
                # Generate answer with insurance expertise
                system_prompt = """You are an expert insurance policy analyst with deep knowledge of Indian insurance regulations and terminology.

Your expertise includes:
- Policy terms, conditions, and exclusions
- Premium calculations and payment terms
- Claim procedures and settlement processes
- Waiting periods and grace periods  
- Coverage limits and deductibles
- Pre-existing disease clauses
- Maternity and specific treatment coverage

Instructions:
1. Provide precise, factual answers based only on the document context
2. Include specific amounts, percentages, and time periods when mentioned
3. Clearly state if information is subject to conditions or exclusions
4. Use proper insurance terminology
5. If information is incomplete, specify what details are missing"""
                
                messages = [
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=f"Document Context (ranked by relevance):\n{context}\n\nQuestion: {question}\n\nPlease provide a comprehensive answer focusing on the specific insurance terms and conditions mentioned in the document.")
                ]
                
                response = await llm.ainvoke(messages)
                return response.content
                
            except Exception as e:
                logger.error(f"Error processing question '{question}': {str(e)}")
                return f"Error processing question: {str(e)}"
        
        # Execute all questions in parallel for maximum speed
        answers = await asyncio.gather(*[process_single_question(q) for q in request.questions])
        
        # Calculate total processing time
        total_time = time.time() - start_time
        logger.info(f"🚀 Completed processing {len(request.questions)} questions across {len(documents_urls)} documents in {total_time:.2f} seconds!")
        
        # Schedule cleanup of embeddings in background
        background_tasks.add_task(cleanup_document_embeddings, documents_urls)
        
        return QueryResponse(answers=answers)
        
    except Exception as e:
        logger.error(f"Error in optimized query retrieval: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

async def cleanup_document_embeddings(doc_urls: List[str]):
    """Background task to clean up document embeddings after processing"""
    try:
        logger.info(f"🧹 Starting cleanup for {len(doc_urls)} documents")
        await pinecone_vector_store.delete_multiple_documents(doc_urls)
        logger.info("✅ Document embeddings cleanup completed")
    except Exception as e:
        logger.error(f"❌ Error during cleanup: {str(e)}")

@app.post("/hackrx/run-legacy", response_model=QueryResponse)
async def run_query_retrieval_legacy(
    request: QueryRequest,
    background_tasks: BackgroundTasks,
    token: str = Depends(verify_token)
):
    """Legacy enhanced adaptive RAG endpoint"""
    start_time = time.time()
    
    try:
        answers = []
        
        for question in request.questions:
            # Initialize enhanced state
            initial_state = {
                "messages": [HumanMessage(content=json.dumps({"documents": request.documents, "query": question}))],
                "query": question,
                "question_type": "general",
                "key_entities": [],
                "documents": [],
                "vector_store": pinecone_vector_store,
                "retrieved_docs": [],
                "entities": {},  # Added for insurance entities
                "answer": "",
                "confidence": 0.0,
                "method_used": "",
                "num_chunks_used": 0
            }
            
            # Run enhanced multi-agent system
            result = agent_graph.invoke(initial_state)
            
            answer = result["answer"]
            method_used = result["method_used"]
            confidence = result["confidence"]
            num_chunks_used = result["num_chunks_used"]
            question_type = result["question_type"]
            
            answers.append(answer)
        
        return QueryResponse(answers=answers)
        
    except Exception as e:
        logger.error(f"Error processing request: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "version": "2.0.0",
        "rag_mode": "adaptive_only"
    }

@app.get("/metrics")
async def get_metrics(token: str = Depends(verify_token)):
    """Basic system metrics without database"""
    return {
        "status": "healthy",
        "pinecone_index": config.PINECONE_INDEX_NAME,
        "embedding_model": config.EMBEDDING_MODEL,
        "version": "2.0.0",
        "features": [
            "pinecone_vector_db",
            "parallel_processing", 
            "insurance_optimization",
            "auto_cleanup"
        ]
    }

@app.post("/webhook")
async def hackathon_webhook(request: dict):
    """Webhook endpoint for hackathon evaluation system"""
    try:
        timestamp = datetime.utcnow().isoformat()
        logger.info(f"🎯 Hackathon webhook received at {timestamp}")
        logger.info(f"Webhook data: {request}")
        
        # System status check
        health_status = await health_check()
        
        response = {
            "status": "success",
            "timestamp": timestamp,
            "system_health": health_status,
            "api_endpoints": {
                "main_submission": "/hackrx/run",
                "health_check": "/health",
                "system_metrics": "/metrics",
                "webhook_status": "/webhook/status"
            },
            "system_capabilities": [
                "multi_document_pdf_processing",
                "parallel_question_processing", 
                "insurance_domain_optimization",
                "pinecone_vector_database",
                "automatic_embedding_cleanup"
            ],
            "performance_stats": {
                "avg_response_time": "8-12 seconds",
                "concurrent_questions": "10+ questions parallel",
                "document_types": ["PDF", "DOCX", "Email"],
                "max_cleanup_time": "2-3 seconds"
            },
            "hackathon_ready": True,
            "webhook_payload_received": request
        }
        
        return response
        
    except Exception as e:
        logger.error(f"❌ Webhook error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/webhook/evaluate")
async def hackathon_evaluation_webhook(request: dict):
    """Webhook for hackathon evaluation results"""
    try:
        logger.info(f"📊 Evaluation webhook received: {request}")
        
        evaluation_response = {
            "status": "evaluation_received",
            "timestamp": datetime.utcnow().isoformat(),
            "evaluation_data": request,
            "system_response": "Ready for next evaluation"
        }
        
        return evaluation_response
        
    except Exception as e:
        logger.error(f"❌ Evaluation webhook error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/webhook/status")
async def webhook_status():
    """Webhook status endpoint for hackathon system"""
    return {
        "webhook_status": "active",
        "api_endpoint": "/hackrx/run",
        "health_endpoint": "/health",
        "metrics_endpoint": "/metrics",
        "last_updated": datetime.utcnow().isoformat(),
        "ready_for_evaluation": True
    }

# Main function
if __name__ == "__main__":
    import uvicorn

    print("""


    
    API Endpoints:
    - POST /hackrx/run - ⚡ ULTRA-FAST Pinecone-optimized endpoint
    - POST /hackrx/run-legacy - Legacy endpoint (fallback)
    - GET /health - Health check
    - GET /metrics - System metrics
    
    Authentication:
    Bearer Token: dbbdb701cfc45d4041e22a03edbfc65753fe9d7b4b9ba1df4884e864f3bb934d


    """)
    
    # Get port from environment variable (Railway sets this automatically)
    port = int(os.environ.get("PORT", 8000))
    
    uvicorn.run(
        app,
        host="0.0.0.0",
        port=port,
        log_level="info"
    )